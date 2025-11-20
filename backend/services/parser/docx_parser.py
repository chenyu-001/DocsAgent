"""DOCX document parser"""
from docx import Document
from utils.text_clean import clean_text


class DOCXParser:
    """Parser for Word documents"""

    def parse(self, file_path: str):
        """Extract text and metadata from DOCX file"""
        doc = Document(file_path)
        text_parts = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        full_text = "\n".join(text_parts)
        cleaned_text = clean_text(full_text)

        # Extract metadata
        core_props = doc.core_properties
        metadata = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
        }

        return {
            "text": cleaned_text,
            "metadata": metadata,
            "page_count": None,
            "word_count": len(cleaned_text.split()),
        }
